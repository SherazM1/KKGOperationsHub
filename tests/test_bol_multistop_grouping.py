from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
import pandas as pd
import pytest

from app.models.bol_multistop_row import BolMultistopRow
from app.services.bol_file_bundle_service import create_multistop_docx_bundle
from app.services.bol_multistop_docx_generator import (
    MultistopGeneratedDocxFile,
    generate_multistop_docx_set,
)
from app.services.bol_multistop_mapper import map_multistop_rows_to_records
from app.services.bol_multistop_parser import (
    _parse_multistop_dataframe_rows,
    _resolve_columns,
    _resolve_optional_columns,
)
from app.utils.bol_facilities import BOL_FACILITY_LOOKUP


def _row(
    *,
    kk_load: str,
    stop: int,
    bol_number: str,
    load_number: str = "",
    source_row_number: int | None = None,
) -> BolMultistopRow:
    return BolMultistopRow(
        source_row_number=source_row_number or stop + 1,
        kk_load=kk_load,
        stop=str(stop),
        stop_number=stop,
        trackers=f"TRK-{kk_load}-{stop}",
        carrier="Carrier",
        load_number=load_number,
        kk_po_number=f"KKPO-{kk_load}",
        bol_number=bol_number,
        ship_date="2026-05-13",
        dc_name=f"DC {stop}",
        dc_address=f"{stop} Test St",
        dc_city_state_zip=f"City {stop}, TX 7500{stop}",
        dc_city=f"City {stop}",
        dc_state="TX",
        dc_zip=f"7500{stop}",
        dc_number=f"DC{stop}",
        target_po_number=f"TGT-{stop}",
        item_number=f"ITEM-{stop}",
        upc=f"UPC-{stop}",
        pallet_description=f"Pallet {stop}",
        cases="10",
        total_pallets="2",
        kit_value_each="",
        shipment_value="",
        chargeback_3_percent="",
        weight_each="100",
        weight="100",
    )


def _fake_multistop_saves(monkeypatch: pytest.MonkeyPatch) -> list[tuple[str, list[int]]]:
    calls: list[tuple[str, list[int]]] = []

    def fake_combined(**kwargs):
        record = kwargs["record"]
        destination = Path(kwargs["output_root"]) / f"{kwargs['base_name']}.docx"
        destination.write_bytes(b"combined")
        calls.append(("combined", [stop.stop_number for stop in record.stops]))
        return MultistopGeneratedDocxFile(
            bol_number=kwargs["bol_label"],
            file_name=destination.name,
            file_path=str(destination),
            document_type="combined",
            load_number=record.load_number,
            kk_load_number=record.kk_load_number,
            template_mode=kwargs.get("template_mode", "Standard"),
            stop_number=None,
        )

    def fake_stop(**kwargs):
        record = kwargs["record"]
        stop = record.stops[kwargs["stop_index"]]
        destination = Path(kwargs["output_root"]) / f"{kwargs['base_name']}.docx"
        destination.write_bytes(b"stop")
        calls.append(("stop", [stop.stop_number]))
        return MultistopGeneratedDocxFile(
            bol_number=stop.bol_number,
            file_name=destination.name,
            file_path=str(destination),
            document_type="stop",
            load_number=record.load_number,
            kk_load_number=record.kk_load_number,
            template_mode=kwargs.get("template_mode", "Standard"),
            stop_number=stop.stop_number,
        )

    monkeypatch.setattr(
        "app.services.bol_multistop_docx_generator._save_multistop_docx",
        fake_combined,
    )
    monkeypatch.setattr(
        "app.services.bol_multistop_docx_generator._save_individual_stop_docx",
        fake_stop,
    )
    return calls


def _capture_multistop_shippers(monkeypatch: pytest.MonkeyPatch) -> dict[str, list[dict[str, str]]]:
    captured: dict[str, list[dict[str, str]]] = {"combined": [], "stop": []}

    def fake_combined(**kwargs):
        record = kwargs["record"]
        selected_facility = kwargs["selected_facility"]
        destination = Path(kwargs["output_root"]) / f"{kwargs['base_name']}.docx"
        destination.write_bytes(b"combined")
        captured["combined"].append(
            {
                "facility_name": selected_facility["facility_name"],
                "address": selected_facility["address"],
                "location": selected_facility["location"],
                "record_company": record.ship_from.company,
                "record_street": record.ship_from.street,
                "record_city_state_zip": record.ship_from.city_state_zip,
                "template_mode": kwargs.get("template_mode", ""),
            }
        )
        return MultistopGeneratedDocxFile(
            bol_number=kwargs["bol_label"],
            file_name=destination.name,
            file_path=str(destination),
            document_type="combined",
            load_number=record.load_number,
            kk_load_number=record.kk_load_number,
            template_mode=kwargs.get("template_mode", "Standard"),
            stop_number=None,
        )

    def fake_apply_template(doc, stop_record, selected_facility, batch_comment, **kwargs):
        captured["stop"].append(
            {
                "facility_name": selected_facility["facility_name"],
                "address": selected_facility["address"],
                "location": selected_facility["location"],
                "record_company": stop_record.ship_from.company,
                "record_street": stop_record.ship_from.street,
                "record_city_state_zip": stop_record.ship_from.city_state_zip,
            }
        )
        return []

    def fake_postprocess(path, resolved_comment):
        return True

    monkeypatch.setattr(
        "app.services.bol_multistop_docx_generator._save_multistop_docx",
        fake_combined,
    )
    monkeypatch.setattr(
        "app.services.bol_multistop_docx_generator._apply_standard_template_record_values",
        fake_apply_template,
    )
    monkeypatch.setattr(
        "app.services.bol_multistop_docx_generator._clean_standard_individual_stop_item_area",
        lambda *args, **kwargs: None,
    )
    monkeypatch.setattr(
        "app.services.bol_multistop_docx_generator._clean_no_recourse_individual_stop_item_area",
        lambda *args, **kwargs: None,
    )
    monkeypatch.setattr(
        "app.services.bol_multistop_docx_generator._postprocess_standard_comments_in_saved_docx",
        fake_postprocess,
    )
    return captured


def _three_stop_record():
    return map_multistop_rows_to_records(
        [
            _row(kk_load="1", stop=1, bol_number="A"),
            _row(kk_load="1", stop=2, bol_number="B"),
            _row(kk_load="1", stop=3, bol_number="C"),
        ]
    )


def _docx_text(path: str) -> str:
    doc = Document(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _assert_captured_shipper(
    captured_values: dict[str, str],
    *,
    facility_name: str,
    street: str,
    city_state_zip: str,
) -> None:
    assert captured_values["facility_name"] == facility_name
    assert captured_values["record_company"] == facility_name
    assert captured_values["record_street"] == street
    assert captured_values["record_city_state_zip"] == city_state_zip


def test_multistop_groups_three_stops_by_kk_load_and_generates_one_combined(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    records = map_multistop_rows_to_records(
        [
            _row(kk_load="1", stop=1, bol_number="A", load_number="L1"),
            _row(kk_load="1", stop=2, bol_number="B", load_number="L2"),
            _row(kk_load="1", stop=3, bol_number="C", load_number=""),
        ]
    )
    calls = _fake_multistop_saves(monkeypatch)

    result = generate_multistop_docx_set(
        records,
        selected_facility={"facility_name": "Test", "address": "A", "location": "B"},
        output_dir=tmp_path / "docx",
    )

    assert len(records) == 1
    assert records[0].kk_load_number == "1"
    assert records[0].group_key == "kk_load::1"
    assert [stop.stop_number for stop in records[0].stops] == [1, 2, 3]
    assert [stop.bol_number for stop in records[0].stops] == ["A", "B", "C"]
    assert len([file for file in result.generated_files if file.document_type == "stop"]) == 3
    assert len([file for file in result.generated_files if file.document_type == "combined"]) == 1
    assert calls[0] == ("combined", [1, 2, 3])

    bundle = create_multistop_docx_bundle(result.generated_files, output_dir=tmp_path / "bundle")
    assert bundle.docx_bundle is not None
    with ZipFile(bundle.docx_bundle.file_path) as zip_file:
        names = sorted(zip_file.namelist())

    assert names == [
        "KK_Load_1/combined_multistop_bol_KK_Load_1.docx",
        "KK_Load_1/stop_1_bol_A.docx",
        "KK_Load_1/stop_2_bol_B.docx",
        "KK_Load_1/stop_3_bol_C.docx",
    ]


def test_multistop_green_bay_shipper_propagates_to_combined_docx_data(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    selected_facility = BOL_FACILITY_LOOKUP["Green Bay Packaging"]
    captured = _capture_multistop_shippers(monkeypatch)

    generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=selected_facility,
        output_dir=tmp_path,
        master_template_mode="Standard",
    )

    assert len(captured["combined"]) == 1
    _assert_captured_shipper(
        captured["combined"][0],
        facility_name="Kendal King C/O Green Bay",
        street="5600 S. Moorland Road",
        city_state_zip="New Berlin, WI 53151",
    )
    assert captured["combined"][0]["record_company"] != "KENDAL KING C/O"


def test_multistop_green_bay_shipper_propagates_to_all_stop_docs(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    selected_facility = BOL_FACILITY_LOOKUP["Green Bay Packaging"]
    captured = _capture_multistop_shippers(monkeypatch)

    generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=selected_facility,
        individual_stop_template_path=Path("app/templates/standard_bol_template.docx"),
        output_dir=tmp_path,
    )

    assert len(captured["stop"]) == 3
    for stop_capture in captured["stop"]:
        _assert_captured_shipper(
            stop_capture,
            facility_name="Kendal King C/O Green Bay",
            street="5600 S. Moorland Road",
            city_state_zip="New Berlin, WI 53151",
        )


def test_multistop_selected_shipper_does_not_fall_back_to_shorr(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    selected_facility = BOL_FACILITY_LOOKUP["Green Bay Packaging"]
    captured = _capture_multistop_shippers(monkeypatch)

    generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=selected_facility,
        output_dir=tmp_path,
    )

    all_captures = [*captured["combined"], *captured["stop"]]
    assert all("Shorr" not in capture["record_company"] for capture in all_captures)
    assert all("981 W Oakdale Rd" not in capture["record_street"] for capture in all_captures)


def test_multistop_another_shipper_selection_propagates_to_all_docs(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    selected_facility = BOL_FACILITY_LOOKUP["SHORR"]
    captured = _capture_multistop_shippers(monkeypatch)

    generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=selected_facility,
        output_dir=tmp_path,
    )

    for capture in [*captured["combined"], *captured["stop"]]:
        _assert_captured_shipper(
            capture,
            facility_name="Kendal King C/O Shorr",
            street="975 W Oakdale Road",
            city_state_zip="Grand Prairie, TX 75050",
        )


def test_multistop_full_selected_shipper_record_propagates_to_all_docs(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    selected_facility = {
        "facility": "TEST",
        "facility_name": "KENDAL KING C/O TEST LOCATION",
        "address": "123 Test Street Test City, WI 54321",
        "location": "Test City, WI",
    }
    captured = _capture_multistop_shippers(monkeypatch)

    generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=selected_facility,
        output_dir=tmp_path,
    )

    for capture in [*captured["combined"], *captured["stop"]]:
        assert capture["facility_name"] == "KENDAL KING C/O TEST LOCATION"
        assert capture["address"] == "123 Test Street Test City, WI 54321"
        assert capture["location"] == "Test City, WI"
        assert capture["record_company"] == "KENDAL KING C/O TEST LOCATION"
        assert capture["record_street"] == "123 Test Street"
        assert capture["record_city_state_zip"] == "Test City, WI 54321"


def test_multistop_combined_docx_renders_full_green_bay_shipper_address(
    tmp_path: Path,
) -> None:
    result = generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=BOL_FACILITY_LOOKUP["Green Bay Packaging"],
        output_dir=tmp_path,
    )

    combined_file = next(file for file in result.generated_files if file.document_type == "combined")
    text = _docx_text(combined_file.file_path)

    assert "Kendal King C/O Green Bay" in text
    assert "5600 S. Moorland Road" in text
    assert "New Berlin, WI 53151" in text


def test_multistop_stop_docx_renders_full_green_bay_shipper_address(
    tmp_path: Path,
) -> None:
    result = generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=BOL_FACILITY_LOOKUP["Green Bay Packaging"],
        individual_stop_template_path=Path("app/templates/standard_bol_template.docx"),
        output_dir=tmp_path,
    )

    stop_files = [file for file in result.generated_files if file.document_type == "stop"]
    assert len(stop_files) == 3
    for stop_file in stop_files:
        text = _docx_text(stop_file.file_path)
        assert "Kendal King C/O Green Bay" in text
        assert "5600 S. Moorland Road" in text
        assert "New Berlin, WI 53151" in text


def test_multistop_generic_shipper_docx_renders_full_address(tmp_path: Path) -> None:
    selected_facility = {
        "facility": "TEST",
        "facility_name": "Kendal King C/O Test Facility",
        "address": "123 Example Ave Madison, WI 53703",
        "location": "Madison, WI",
    }

    result = generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=selected_facility,
        output_dir=tmp_path,
    )

    for generated_file in result.generated_files:
        text = _docx_text(generated_file.file_path)
        assert "Kendal King C/O Test Facility" in text
        assert "123 Example Ave" in text
        assert "Madison, WI 53703" in text


def test_multistop_standard_mode_propagates_to_master_and_stops(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    captured = _capture_multistop_shippers(monkeypatch)

    result = generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=BOL_FACILITY_LOOKUP["Green Bay Packaging"],
        individual_stop_template_path=Path("app/templates/standard_bol_template.docx"),
        master_template_mode="Standard",
        output_dir=tmp_path,
    )

    assert [file.template_mode for file in result.generated_files] == [
        "Standard",
        "Standard",
        "Standard",
        "Standard",
    ]
    assert captured["combined"][0]["template_mode"] == "Standard"


def test_multistop_no_recourse_mode_propagates_to_master_and_stops(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    captured = _capture_multistop_shippers(monkeypatch)

    result = generate_multistop_docx_set(
        _three_stop_record(),
        selected_facility=BOL_FACILITY_LOOKUP["Green Bay Packaging"],
        individual_stop_template_path=Path("app/templates/no_recourse_bol_template.docx"),
        master_template_mode="No Recourse",
        output_dir=tmp_path,
    )

    assert [file.template_mode for file in result.generated_files] == [
        "No Recourse",
        "No Recourse",
        "No Recourse",
        "No Recourse",
    ]
    assert captured["combined"][0]["template_mode"] == "No Recourse"


def test_multistop_groups_multiple_loads_without_cross_contamination() -> None:
    records = map_multistop_rows_to_records(
        [
            _row(kk_load="1", stop=1, bol_number="A"),
            _row(kk_load="1", stop=2, bol_number="B"),
            _row(kk_load="1", stop=3, bol_number="C"),
            _row(kk_load="2", stop=1, bol_number="D"),
            _row(kk_load="2", stop=2, bol_number="E"),
        ]
    )

    assert len(records) == 2
    by_load = {record.kk_load_number: record for record in records}
    assert [stop.stop_number for stop in by_load["1"].stops] == [1, 2, 3]
    assert [stop.bol_number for stop in by_load["1"].stops] == ["A", "B", "C"]
    assert [stop.stop_number for stop in by_load["2"].stops] == [1, 2]
    assert [stop.bol_number for stop in by_load["2"].stops] == ["D", "E"]


def test_multistop_unique_bol_numbers_do_not_split_kk_load_group() -> None:
    records = map_multistop_rows_to_records(
        [
            _row(kk_load="1", stop=1, bol_number="A"),
            _row(kk_load="1", stop=2, bol_number="B"),
            _row(kk_load="1", stop=3, bol_number="C"),
        ]
    )

    assert len(records) == 1
    assert [stop.bol_number for stop in records[0].stops] == ["A", "B", "C"]


def test_multistop_load_number_column_does_not_control_grouping() -> None:
    records = map_multistop_rows_to_records(
        [
            _row(kk_load="1", stop=1, bol_number="A", load_number=""),
            _row(kk_load="1", stop=2, bol_number="B", load_number="DIFFERENT"),
            _row(kk_load="1", stop=3, bol_number="C", load_number="OTHER"),
        ]
    )

    assert len(records) == 1
    assert records[0].kk_load_number == "1"
    assert records[0].load_number == "DIFFERENT"
    assert [stop.stop_number for stop in records[0].stops] == [1, 2, 3]


def test_multistop_stops_are_ordered_numerically_within_kk_load() -> None:
    records = map_multistop_rows_to_records(
        [
            _row(kk_load="1", stop=3, bol_number="C", source_row_number=2),
            _row(kk_load="1", stop=1, bol_number="A", source_row_number=3),
            _row(kk_load="1", stop=2, bol_number="B", source_row_number=4),
        ]
    )

    assert [stop.stop_number for stop in records[0].stops] == [1, 2, 3]
    assert [stop.bol_number for stop in records[0].stops] == ["A", "B", "C"]


def test_halloween_workbook_multistop_grouping_by_kk_load_if_available() -> None:
    workbook_path = Path(r"C:\Users\shera\Downloads\Load_Sheet_halloween_2026_burts.xlsx")
    if not workbook_path.exists():
        pytest.skip("Halloween Multistop load sheet fixture is not available locally.")

    workbook = pd.ExcelFile(workbook_path)
    sheet_name = workbook.sheet_names[0]
    df = workbook.parse(sheet_name=sheet_name, dtype=object)
    column_map = _resolve_columns(df.columns.tolist(), sheet_name)
    optional_column_map = _resolve_optional_columns(df.columns.tolist())
    rows = _parse_multistop_dataframe_rows(df, column_map, optional_column_map)
    records = map_multistop_rows_to_records(rows)

    assert len(records) == 12
    by_load = {record.kk_load_number: record for record in records}
    for load_number in ("1", "5", "6", "7", "11", "12"):
        assert len(by_load[load_number].stops) == 3
