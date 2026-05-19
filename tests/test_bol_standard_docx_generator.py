from __future__ import annotations

from pathlib import Path

from docx import Document

from app.models.bol_standard_record import (
    BolAddressBlock,
    BolStandardItemLine,
    BolStandardRecord,
)
from app.services.bol_standard_docx_generator import (
    generate_standard_docx_set,
    resolve_output_filename_prefix_for_mode,
    resolve_template_path_for_mode,
)
from app.utils.bol_facilities import BOL_FACILITY_LOOKUP, BOL_FACILITY_OPTIONS
from app.utils.bol_facilities import facility_to_ship_from


def _ready_record() -> BolStandardRecord:
    return BolStandardRecord(
        bol_number="10001859231-0553",
        ship_date="2026-05-13",
        carrier="Test Carrier",
        kk_load_number="1",
        kk_po_number="KKPO-001",
        po_number="10001859231-0553",
        dc_number="0553",
        consignee_company="Test DC",
        consignee_street="123 Test Street",
        consignee_city_state_zip="Dallas, TX 75001",
        ship_from=BolAddressBlock(
            company="Kendal King C/O Shorr",
            street="981 W Oakdale Rd",
            city_state_zip="Grand Prairie, TX 75056",
        ),
        bill_to=BolAddressBlock(
            company="Trident Transport, LLC",
            street="505 Riverfront Pkwy",
            city_state_zip="Chattanooga, TN 37402",
        ),
        seal_number_blank="",
        comments="",
        item_lines=[
            BolStandardItemLine(
                source_row_number=2,
                pallet_qty="2",
                type="PLT",
                po_number="10001859231-0553",
                item_description="Test pallet",
                item_number="ITEM1",
                upc="000111222333",
                skids="2",
                weight_each="100",
                total_weight="306 lbs.",
            )
        ],
        total_skids=2,
        is_ready=True,
        status="Ready",
        carrier_pro_number="1073839",
        pickup_number="PU-123",
    )


def _generated_docx(mode: str, tmp_path: Path, *, bol_type: str, qty_type: str) -> Document:
    result = generate_standard_docx_set(
        [_ready_record()],
        selected_facility=BOL_FACILITY_LOOKUP[BOL_FACILITY_OPTIONS[0]],
        bol_type=bol_type,
        qty_type=qty_type,
        template_path=resolve_template_path_for_mode(mode),
        output_dir=tmp_path,
        file_name_prefix=resolve_output_filename_prefix_for_mode(mode),
    )

    assert result.generated_count == 1
    return Document(result.generated_files[0].file_path)


def _item_header_text(doc: Document) -> str:
    table = doc.tables[0]
    for row in table.rows:
        row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
        row_text_upper = row_text.upper()
        if "TYPE" in row_text_upper and "PO #" in row_text_upper and "ITEM DESCRIPTION" in row_text_upper:
            return row.cells[0].text.strip()
    raise AssertionError("Item table header row was not found.")


def _first_item_type_value(doc: Document) -> str:
    table = doc.tables[0]
    for row in table.rows:
        row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
        if "Test pallet" in row_text:
            return row.cells[1].text.strip()
    raise AssertionError("First item row was not found.")


def _document_text(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _totals_row_text(doc: Document) -> str:
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if "TOTALS" in row_text.upper():
                return row_text
    raise AssertionError("Totals row was not found.")


def test_standard_docx_qty_type_plt_renders_pallet_qty_header(tmp_path: Path) -> None:
    doc = _generated_docx("Standard", tmp_path, bol_type="CASE", qty_type="PLT")

    assert _item_header_text(doc) == "Pallet Qty"
    assert _first_item_type_value(doc) == "CASE"


def test_standard_docx_qty_type_case_renders_case_qty_header(tmp_path: Path) -> None:
    doc = _generated_docx("Standard", tmp_path, bol_type="PLT", qty_type="Case")

    assert _item_header_text(doc) == "Case Qty"
    assert _first_item_type_value(doc) == "PLT"


def test_no_recourse_docx_qty_type_case_renders_case_qty_header(tmp_path: Path) -> None:
    doc = _generated_docx("No Recourse", tmp_path, bol_type="PLT", qty_type="Case")

    assert _item_header_text(doc) == "Case Qty"
    assert _first_item_type_value(doc) == "PLT"


def test_standard_docx_totals_prefer_total_weight_and_pickup_renders(tmp_path: Path) -> None:
    doc = _generated_docx("Standard", tmp_path, bol_type="PLT", qty_type="PLT")

    assert "PU-123" in _document_text(doc)
    assert "306" in _totals_row_text(doc)


def test_standard_docx_can_suppress_pickup_number_without_changing_record(tmp_path: Path) -> None:
    record = _ready_record()

    result = generate_standard_docx_set(
        [record],
        selected_facility=BOL_FACILITY_LOOKUP[BOL_FACILITY_OPTIONS[0]],
        bol_type="PLT",
        qty_type="PLT",
        render_pickup_number=False,
        template_path=resolve_template_path_for_mode("Standard"),
        output_dir=tmp_path,
        file_name_prefix=resolve_output_filename_prefix_for_mode("Standard"),
    )
    doc = Document(result.generated_files[0].file_path)

    assert record.pickup_number == "PU-123"
    assert "PU-123" not in _document_text(doc)
    assert "306" in _totals_row_text(doc)


def test_standard_docx_renders_carrier_pro_from_load_number_not_kk_load(tmp_path: Path) -> None:
    doc = _generated_docx("Standard", tmp_path, bol_type="PLT", qty_type="PLT")
    text = _document_text(doc)

    assert "1073839" in text
    assert "1" in text


def test_standard_docx_type_case_has_no_inserted_spaces_or_line_breaks(tmp_path: Path) -> None:
    doc = _generated_docx("Standard", tmp_path, bol_type="CASE", qty_type="PLT")
    type_value = _first_item_type_value(doc)

    assert type_value == "CASE"


def test_standard_docx_uses_record_ship_from_from_selected_facility(tmp_path: Path) -> None:
    record = _ready_record()
    selected_facility = BOL_FACILITY_LOOKUP["PRODUCTIV-ESTERS"]
    record.ship_from = facility_to_ship_from(selected_facility)

    result = generate_standard_docx_set(
        [record],
        selected_facility=selected_facility,
        bol_type="PLT",
        qty_type="PLT",
        template_path=resolve_template_path_for_mode("Standard"),
        output_dir=tmp_path,
        file_name_prefix=resolve_output_filename_prefix_for_mode("Standard"),
    )
    doc = Document(result.generated_files[0].file_path)
    text = _document_text(doc)

    assert "Kendal King C/O Productiv" in text
    assert "2450 Esters BLVD Suite 100" in text
    assert "C A S E" not in _document_text(doc)
    assert "CAS\nE" not in _document_text(doc)


def test_no_recourse_docx_type_case_has_no_inserted_spaces_or_line_breaks(tmp_path: Path) -> None:
    doc = _generated_docx("No Recourse", tmp_path, bol_type="CASE", qty_type="PLT")
    type_value = _first_item_type_value(doc)

    assert type_value == "CASE"
    assert "C A S E" not in _document_text(doc)
    assert "CAS\nE" not in _document_text(doc)
