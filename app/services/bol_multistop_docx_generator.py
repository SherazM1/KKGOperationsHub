"""DOCX generation service for Multistop-mode BOL records."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from tempfile import mkdtemp

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from docx.table import Table

from app.models.bol_multistop_record import BolMultistopRecord
from app.models.bol_standard_record import BolStandardItemLine, BolStandardRecord
from app.services.bol_standard_docx_generator import (
    DocxGenerationNotice,
    FailedDocxRecord,
    GeneratedDocxFile,
    NO_RECOURSE_TEMPLATE_PATH,
    STANDARD_TEMPLATE_PATH,
    SkippedDocxRecord,
    StandardDocxGenerationResult,
    _apply_template_record_values as _apply_standard_template_record_values,
    _postprocess_comments_in_saved_docx as _postprocess_standard_comments_in_saved_docx,
)
from app.utils.bol_facilities import BolFacilityRecord


MULTISTOP_TEMPLATE_PATH = Path("app/templates/multistop_bol_template.docx")
LEFT_MERGE = "\u00ab"
RIGHT_MERGE = "\u00bb"


@dataclass(slots=True)
class MultistopGeneratedDocxFile(GeneratedDocxFile):
    """Generated Multistop DOCX metadata for combined and stop-level outputs."""

    document_type: str
    load_number: str
    stop_number: int | None = None


def _tok(name: str) -> str:
    return f"{LEFT_MERGE}{name}{RIGHT_MERGE}"


def _sanitize_filename_part(value: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in ("-", "_") else "_" for char in value)
    cleaned = cleaned.strip("_")
    return cleaned or "unknown"


def _unique_destination_path(directory: Path, base_name: str, extension: str) -> Path:
    candidate = directory / f"{base_name}{extension}"
    if not candidate.exists():
        return candidate

    suffix = 2
    while True:
        candidate = directory / f"{base_name}_{suffix}{extension}"
        if not candidate.exists():
            return candidate
        suffix += 1


def _format_number(value: float) -> str:
    return str(int(value)) if float(value).is_integer() else f"{value:.2f}".rstrip("0").rstrip(".")


def _parse_number(value: str) -> float:
    cleaned = (value or "").replace(",", "").strip()
    if not cleaned:
        return 0.0
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def _format_ship_date_for_template(raw_ship_date: str) -> str:
    value = (raw_ship_date or "").strip()
    if not value:
        return ""

    normalized = value.replace("T", " ")
    for sep in (" ", "."):
        if sep in normalized:
            date_candidate = normalized.split(sep, 1)[0].strip()
            if date_candidate and any(char.isdigit() for char in date_candidate):
                normalized = date_candidate
                break

    parse_formats = (
        "%Y-%m-%d",
        "%m/%d/%Y",
        "%m/%d/%y",
        "%m-%d-%Y",
        "%m-%d-%y",
    )
    for parse_format in parse_formats:
        try:
            parsed = datetime.strptime(normalized, parse_format)
            return parsed.strftime("%m/%d/%Y")
        except ValueError:
            continue

    try:
        parsed_iso = datetime.fromisoformat(value.replace("Z", ""))
        return parsed_iso.strftime("%m/%d/%Y")
    except ValueError:
        return normalized


def _set_text_node_value(node, value: str) -> None:
    if "\n" not in value:
        node.text = value
        return

    run = node.getparent()
    insert_at = run.index(node)
    run.remove(node)

    parts = value.split("\n")
    for index, part in enumerate(parts):
        if index > 0:
            br = OxmlElement("w:br")
            run.insert(insert_at, br)
            insert_at += 1

        text_node = OxmlElement("w:t")
        if part.startswith(" ") or part.endswith(" "):
            text_node.set(qn("xml:space"), "preserve")
        text_node.text = part
        run.insert(insert_at, text_node)
        insert_at += 1


def _replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    text_nodes = paragraph._p.findall(".//w:t", paragraph._p.nsmap)
    instr_nodes = paragraph._p.findall(".//w:instrText", paragraph._p.nsmap)
    for node in [*text_nodes, *instr_nodes]:
        text = node.text or ""
        updated = text
        for source, target in replacements.items():
            if source in updated:
                updated = updated.replace(source, target)
        if updated != text:
            _set_text_node_value(node, updated)


def _replace_text_in_document(
    doc: Document, replacements: dict[str, str], *, include_xml_tree: bool = True
) -> None:
    def _replace_in_table_collection(tables: list[Table]) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, replacements)
                    _replace_in_table_collection(cell.tables)

    def _replace_in_element_tree(element) -> None:
        text_nodes = element.findall(".//w:t", element.nsmap)
        instr_nodes = element.findall(".//w:instrText", element.nsmap)
        for node in [*text_nodes, *instr_nodes]:
            text = node.text or ""
            updated = text
            for source, target in replacements.items():
                if source in updated:
                    updated = updated.replace(source, target)
            if updated != text:
                _set_text_node_value(node, updated)

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    _replace_in_table_collection(doc.tables)
    if include_xml_tree:
        _replace_in_element_tree(doc.element)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        _replace_in_table_collection(section.header.tables)
        if include_xml_tree:
            _replace_in_element_tree(section.header._element)

        for paragraph in section.footer.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        _replace_in_table_collection(section.footer.tables)
        if include_xml_tree:
            _replace_in_element_tree(section.footer._element)


def _set_row_height(row, twips: int, *, exact: bool = True) -> None:
    row.height = Twips(twips)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST


def _compact_row_text(row, font_points: float = 8.5) -> None:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 0.95
            for run in paragraph.runs:
                run.font.size = Pt(font_points)


def _set_cell_text(cell, value: str, *, font_points: float = 8.5, align_center: bool = True) -> None:
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.95
        for run in paragraph.runs:
            run.font.size = Pt(font_points)


def _format_multistop_item_description(
    description: str,
    item_number: str,
    upc: str,
) -> str:
    description_value = (description or "").strip()
    detail_parts: list[str] = []
    item_number_value = (item_number or "").strip()
    upc_value = (upc or "").strip()

    if item_number_value:
        detail_parts.append(f"Item #: {item_number_value}")
    if upc_value:
        detail_parts.append(f"UPC #: {upc_value}")

    if not detail_parts:
        return description_value

    detail_line = "    ".join(detail_parts)
    if description_value:
        return f"{description_value}\n{detail_line}"
    return detail_line


def _set_unique_cell_text(
    row,
    cell_index: int,
    value: str,
    *,
    font_points: float = 8.5,
    align_center: bool = True,
) -> None:
    if cell_index >= len(row.cells):
        return
    cell = row.cells[cell_index]
    _set_cell_text(cell, value, font_points=font_points, align_center=align_center)


def _set_reference_field_cell_text(
    cell,
    value: str,
    *,
    font_points: float = 9.0,
    align_center: bool = False,
) -> None:
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else None
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_points)
            run.bold = False


def _set_unique_reference_field_text(
    row,
    cell_index: int,
    value: str,
    *,
    font_points: float = 9.0,
    align_center: bool = False,
) -> None:
    if cell_index >= len(row.cells):
        return
    _set_reference_field_cell_text(
        row.cells[cell_index],
        value,
        font_points=font_points,
        align_center=align_center,
    )


def _clear_row_text(row) -> None:
    seen_cells = set()
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id in seen_cells:
            continue
        seen_cells.add(cell_id)
        cell.text = ""


def _row_unique_cells(row) -> list[tuple[int, object]]:
    unique_cells: list[tuple[int, object]] = []
    seen_cells = set()
    for index, cell in enumerate(row.cells):
        cell_id = id(cell._tc)
        if cell_id in seen_cells:
            continue
        seen_cells.add(cell_id)
        unique_cells.append((index, cell))
    return unique_cells


def _cell_header_text(header_row, cell_index: int) -> str:
    if cell_index >= len(header_row.cells):
        return ""
    return header_row.cells[cell_index].text.strip().upper()


def _is_bol_item_detail_header(row) -> bool:
    row_text_upper = " ".join(cell.text.strip() for cell in row.cells).upper()
    return (
        "ITEM DESCRIPTION" in row_text_upper
        and "PO #" in row_text_upper
        and "WEIGHT" in row_text_upper
        and (
            "QTY" in row_text_upper
            or "PALLET QTY" in row_text_upper
            or "CASE" in row_text_upper
        )
    )


def _row_has_visible_text(row) -> bool:
    return any(cell.text.strip() for _, cell in _row_unique_cells(row))


def _is_item_description_cell(header_row, cell_index: int) -> bool:
    return "ITEM DESCRIPTION" in _cell_header_text(header_row, cell_index)


def _format_item_detail_row(
    row,
    header_row,
    *,
    font_points: float,
    bold: bool = False,
) -> None:
    for cell_index, cell in _row_unique_cells(row):
        is_description = _is_item_description_cell(header_row, cell_index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if is_description and not bold else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 0.95
            for run in paragraph.runs:
                run.font.size = Pt(font_points)
                run.bold = bold


def format_bol_item_detail_table(table: Table) -> None:
    header_idx = None
    for index, row in enumerate(table.rows):
        if _is_bol_item_detail_header(row):
            header_idx = index
            break

    if header_idx is None:
        return

    table.autofit = False
    header_row = table.rows[header_idx]
    _set_row_height(header_row, 300, exact=False)
    _format_item_detail_row(header_row, header_row, font_points=9.0, bold=True)

    for row in table.rows[header_idx + 1 :]:
        row_text_upper = " ".join(cell.text.strip() for cell in row.cells).upper()
        if not _row_has_visible_text(row):
            continue
        if "SHIPPER SIGNATURE" in row_text_upper:
            break

        is_totals_row = "TOTALS" in row_text_upper
        _set_row_height(row, 430 if is_totals_row else 520, exact=False)
        _format_item_detail_row(
            row,
            header_row,
            font_points=8.5,
            bold=False,
        )
        if is_totals_row:
            break


def _tighten_multistop_template_rows(doc: Document) -> None:
    for table in doc.tables:
        rows = list(table.rows)
        for index, row in enumerate(rows):
            row_text = " ".join(cell.text for cell in row.cells)

            if any(token in row_text for token in ("DELIVERY_1_DC", "DELIVERY_2_DC", "DELIVERY_3_DC")):
                _set_row_height(row, 285)
                _compact_row_text(row, 7.5)
            elif any(
                token in row_text
                for token in (
                    "DELIVERY_1_ADDRESS",
                    "DELIVERY_2_ADDRESS",
                    "DELIVERY_3_ADDRESS",
                )
            ):
                _set_row_height(row, 365)
                _compact_row_text(row, 8.0)
            elif any(token in row_text for token in ("DC_1", "DC_2", "DC_3")):
                _set_row_height(row, 520, exact=False)
                _compact_row_text(row, 8.5)

            if 27 <= index <= 32 and not row_text.strip():
                _set_row_height(row, 115)
                _compact_row_text(row, 8.0)


def _clean_standard_individual_stop_item_area(doc: Document, stop) -> None:
    for table in doc.tables:
        header_idx = None
        for idx, row in enumerate(table.rows):
            row_text_upper = " ".join(cell.text.strip() for cell in row.cells).upper()
            if (
                "QTY" in row_text_upper
                and "TYPE" in row_text_upper
                and "PO #" in row_text_upper
                and "ITEM DESCRIPTION" in row_text_upper
                and "# SKIDS" in row_text_upper
                and "WEIGHT" in row_text_upper
            ):
                header_idx = idx
                break

        if header_idx is None:
            continue

        item_idx = header_idx + 1
        totals_idx = None
        for idx in range(item_idx + 1, len(table.rows)):
            row_text_upper = " ".join(cell.text.strip() for cell in table.rows[idx].cells).upper()
            if "TOTALS" in row_text_upper:
                totals_idx = idx
                break

        if item_idx >= len(table.rows) or totals_idx is None:
            return

        item_row = table.rows[item_idx]
        _set_row_height(item_row, 520, exact=False)
        _compact_row_text(item_row, 8.5)
        _set_unique_cell_text(item_row, 0, stop.cases, font_points=8.5)
        _set_unique_cell_text(item_row, 1, "PLT", font_points=8.5)
        _set_unique_cell_text(item_row, 3, stop.target_po_number, font_points=8.0)
        _set_unique_cell_text(
            item_row,
            5,
            _format_multistop_item_description(
                stop.pallet_description,
                stop.item_number,
                stop.upc,
            ),
            font_points=8.5,
            align_center=False,
        )
        _set_unique_cell_text(item_row, 11, stop.total_pallets, font_points=8.5)
        _set_unique_cell_text(item_row, 14, stop.weight, font_points=8.5)

        totals_row = table.rows[totals_idx]
        _set_row_height(totals_row, 430, exact=False)
        _compact_row_text(totals_row, 8.5)
        _set_unique_cell_text(totals_row, 0, stop.cases, font_points=8.5)
        _set_unique_cell_text(totals_row, 5, "TOTALS", font_points=8.5)
        _set_unique_cell_text(totals_row, 11, stop.total_pallets, font_points=8.5)
        _set_unique_cell_text(totals_row, 14, stop.weight, font_points=8.5)
        format_bol_item_detail_table(table)
        return


def _clean_no_recourse_individual_stop_item_area(doc: Document, stop) -> None:
    for table in doc.tables:
        header_idx = None
        for idx, row in enumerate(table.rows):
            row_text_upper = " ".join(cell.text.strip() for cell in row.cells).upper()
            if (
                "PALLET QTY" in row_text_upper
                and "TYPE" in row_text_upper
                and "PO #" in row_text_upper
                and "ITEM DESCRIPTION" in row_text_upper
                and "# SKIDS" in row_text_upper
                and "WEIGHT" in row_text_upper
            ):
                header_idx = idx
                break

        if header_idx is None:
            continue

        item_idx = header_idx + 1
        totals_idx = None
        signature_idx = None
        for idx in range(item_idx + 1, len(table.rows)):
            row_text_upper = " ".join(cell.text.strip() for cell in table.rows[idx].cells).upper()
            if totals_idx is None and "TOTALS" in row_text_upper:
                totals_idx = idx
            if "SHIPPER SIGNATURE" in row_text_upper:
                signature_idx = idx
                break

        if item_idx >= len(table.rows) or totals_idx is None:
            return

        item_row = table.rows[item_idx]
        _set_row_height(item_row, 520, exact=False)
        _compact_row_text(item_row, 8.5)
        _set_unique_cell_text(item_row, 0, stop.cases, font_points=8.5)
        _set_unique_cell_text(item_row, 1, "PLT", font_points=8.5)
        _set_unique_cell_text(item_row, 3, stop.target_po_number, font_points=8.0)
        _set_unique_cell_text(
            item_row,
            5,
            _format_multistop_item_description(
                stop.pallet_description,
                stop.item_number,
                stop.upc,
            ),
            font_points=8.5,
            align_center=False,
        )
        _set_unique_cell_text(item_row, 11, stop.total_pallets, font_points=8.5)
        _set_unique_cell_text(item_row, 14, stop.weight, font_points=8.5)

        stop_unused_at = signature_idx if signature_idx is not None else totals_idx
        for idx in range(item_idx + 1, stop_unused_at):
            if idx == totals_idx:
                continue
            row = table.rows[idx]
            _clear_row_text(row)
            _set_row_height(row, 1)
            _compact_row_text(row, 1.0)

        totals_row = table.rows[totals_idx]
        _set_row_height(totals_row, 430, exact=False)
        _compact_row_text(totals_row, 8.5)
        _set_unique_cell_text(totals_row, 0, stop.cases, font_points=8.5)
        _set_unique_cell_text(totals_row, 5, "TOTALS", font_points=8.5)
        _set_unique_cell_text(totals_row, 11, stop.total_pallets, font_points=8.5)
        _set_unique_cell_text(totals_row, 14, stop.weight, font_points=8.5)

        if signature_idx is not None:
            signature_row = table.rows[signature_idx]
            for cell_index in (11, 14):
                _set_unique_cell_text(signature_row, cell_index, "", font_points=8.5)
        format_bol_item_detail_table(table)
        return


def _clean_combined_multistop_item_area(doc: Document, record: BolMultistopRecord) -> None:
    for table in doc.tables:
        header_idx = None
        for idx, row in enumerate(table.rows):
            row_text_upper = " ".join(cell.text.strip() for cell in row.cells).upper()
            if (
                "DC" in row_text_upper
                and "CASE" in row_text_upper
                and "PO #" in row_text_upper
                and "ITEM DESCRIPTION" in row_text_upper
                and "PALLET" in row_text_upper
                and "WEIGHT" in row_text_upper
            ):
                header_idx = idx
                break

        if header_idx is None:
            continue

        totals_idx = None
        for idx in range(header_idx + 1, len(table.rows)):
            row_text_upper = " ".join(cell.text.strip() for cell in table.rows[idx].cells).upper()
            if "TOTALS" in row_text_upper:
                totals_idx = idx
                break

        if totals_idx is None:
            return

        for stop_offset, stop in enumerate(record.stops):
            row_idx = header_idx + 1 + stop_offset
            if row_idx >= totals_idx:
                break
            _set_unique_cell_text(
                table.rows[row_idx],
                6,
                _format_multistop_item_description(
                    stop.pallet_description,
                    stop.item_number,
                    stop.upc,
                ),
                font_points=8.5,
                align_center=False,
            )

        totals_row = table.rows[totals_idx]
        _set_row_height(totals_row, 430, exact=False)
        _compact_row_text(totals_row, 8.5)
        _set_unique_cell_text(totals_row, 1, _format_number(record.total_case), font_points=8.5)
        _set_unique_cell_text(totals_row, 6, "TOTALS", font_points=8.5)
        _set_unique_cell_text(totals_row, 12, _format_number(record.total_pallet), font_points=8.5)
        _set_unique_cell_text(
            totals_row,
            15,
            _format_number(record.total_ship_weight),
            font_points=8.5,
        )
        format_bol_item_detail_table(table)
        return


def _resolve_comment_for_record(record_comment: str, batch_comment: str | None) -> str:
    record_value = (record_comment or "").strip()
    if record_value:
        return record_value
    return (batch_comment or "").strip()


def _populate_ship_from_block(doc: Document, selected_facility: BolFacilityRecord) -> bool:
    name_value = selected_facility["facility_name"]
    street_value = selected_facility["address"]
    city_state_zip_value = selected_facility["location"]

    for table in doc.tables:
        in_ship_from_block = False
        for row in table.rows:
            row_cells = row.cells
            if not row_cells:
                continue

            row_text_upper = " ".join(cell.text.strip() for cell in row_cells).upper()
            first_cell_text = row_cells[0].text.strip().upper().replace(" ", "")

            if "FROM (SHIPPER)" in row_text_upper:
                in_ship_from_block = True
                continue
            if in_ship_from_block and "TO (CONSIGNEE)" in row_text_upper:
                return True
            if not in_ship_from_block:
                continue

            if first_cell_text == "NAME" and len(row_cells) > 1:
                row_cells[1].text = name_value
            elif first_cell_text == "STREET" and len(row_cells) > 1:
                row_cells[1].text = street_value
            elif first_cell_text in {"CITY/ST/ZIP", "CITY/STATE/ZIP"} and len(row_cells) > 1:
                row_cells[1].text = city_state_zip_value

    return False


def _populate_combined_bill_to_block(doc: Document, record: BolMultistopRecord) -> bool:
    bill_to_lines = [
        record.bill_to.company,
        record.bill_to.street,
        record.bill_to.city_state_zip,
        f"Attn: {record.bill_to.attn}".rstrip(),
    ]

    for table in doc.tables:
        bill_to_row_idx = None
        bill_to_cell_idx = None
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in _row_unique_cells(row):
                if "BILL TO:" in cell.text.upper():
                    bill_to_row_idx = row_idx
                    bill_to_cell_idx = cell_idx
                    break
            if bill_to_row_idx is not None:
                break

        if bill_to_row_idx is None or bill_to_cell_idx is None:
            continue

        for offset, line in enumerate(bill_to_lines, start=1):
            target_row_idx = bill_to_row_idx + offset
            if target_row_idx >= len(table.rows):
                break
            _set_unique_reference_field_text(
                table.rows[target_row_idx],
                bill_to_cell_idx,
                line,
                font_points=9.0,
                align_center=False,
            )
        return True

    return False


def _fit_combined_bol_number(doc: Document, bol_number: str) -> bool:
    if not (bol_number or "").strip():
        return False

    for table in doc.tables:
        for row in table.rows:
            for cell_idx, cell in _row_unique_cells(row):
                if cell.text.strip() != bol_number:
                    continue

                bol_length = len(bol_number.strip())
                font_points = 9.0
                if bol_length > 22:
                    font_points = 8.0
                elif bol_length > 16:
                    font_points = 8.5

                _set_unique_reference_field_text(
                    row,
                    cell_idx,
                    bol_number,
                    font_points=font_points,
                    align_center=True,
                )
                return True

    return False


def _template_replacements(record: BolMultistopRecord) -> dict[str, str]:
    replacements = {
        _tok("BOL_"): record.bol_number,
        _tok("ship_date"): _format_ship_date_for_template(record.ship_date),
        _tok("Carrier"): record.carrier,
        _tok("load"): record.load_number,
        _tok("KK_PO"): record.kk_po_number,
        _tok("KK_Load"): record.kk_load_number,
        _tok("DELIVERY_1_DC"): record.delivery_1_dc,
        _tok("DELIVERY_1_ADDRESS"): record.delivery_1_address,
        _tok("DELIVERY_2_DC"): record.delivery_2_dc,
        _tok("DELIVERY_2_ADDRESS"): record.delivery_2_address,
        _tok("DELIVERY_3_DC"): record.delivery_3_dc,
        _tok("DELIVERY_3_ADDRESS"): record.delivery_3_address,
        _tok("DC_1"): record.dc_1,
        _tok("CASE_1"): record.case_1,
        _tok("PO_1"): record.po_1,
        _tok("Pallet_Description_1"): (
            _format_multistop_item_description(
                record.stops[0].pallet_description,
                record.stops[0].item_number,
                record.stops[0].upc,
            )
            if len(record.stops) > 0
            else ""
        ),
        _tok("PLT_1"): record.plt_1,
        _tok("WEIGHT_1"): record.weight_1,
        _tok("DC_2"): record.dc_2,
        _tok("CASE_2"): record.case_2,
        _tok("PO_2"): record.po_2,
        _tok("Pallet_Description_2"): (
            _format_multistop_item_description(
                record.stops[1].pallet_description,
                record.stops[1].item_number,
                record.stops[1].upc,
            )
            if len(record.stops) > 1
            else ""
        ),
        _tok("PLT_2"): record.plt_2,
        _tok("WEIGHT_2"): record.weight_2,
        _tok("DC_3"): record.dc_3,
        _tok("CASE_3"): record.case_3,
        _tok("PO_3"): record.po_3,
        _tok("Pallet_Description_3"): (
            _format_multistop_item_description(
                record.stops[2].pallet_description,
                record.stops[2].item_number,
                record.stops[2].upc,
            )
            if len(record.stops) > 2
            else ""
        ),
        _tok("PLT_3"): record.plt_3,
        _tok("WEIGHT_3"): record.weight_3,
        _tok("Total_Case"): _format_number(record.total_case),
        _tok("Total_Pallet"): _format_number(record.total_pallet),
        _tok("Total_Ship_Weight"): _format_number(record.total_ship_weight),
    }
    return replacements


def _build_individual_stop_standard_record(
    record: BolMultistopRecord,
    stop_index: int,
) -> BolStandardRecord:
    stop = record.stops[stop_index]
    return BolStandardRecord(
        bol_number=record.bol_number,
        ship_date=record.ship_date,
        carrier=record.carrier,
        kk_load_number=record.kk_load_number,
        kk_po_number=record.kk_po_number,
        po_number=stop.target_po_number,
        dc_number=stop.dc_number,
        consignee_company=stop.delivery_dc,
        consignee_street=stop.delivery_address,
        consignee_city_state_zip=stop.delivery_city_state_zip,
        ship_from=record.ship_from,
        bill_to=record.bill_to,
        seal_number_blank="",
        comments=record.comments,
        item_lines=[
            BolStandardItemLine(
                source_row_number=stop.source_row_number,
                pallet_qty=stop.cases,
                type="PLT",
                po_number=stop.target_po_number,
                item_description=_format_multistop_item_description(
                    stop.pallet_description,
                    stop.item_number,
                    stop.upc,
                ),
                item_number=stop.item_number,
                upc=stop.upc,
                skids=stop.total_pallets,
                weight_each=stop.weight,
            )
        ],
        total_skids=_parse_number(stop.cases),
        is_ready=True,
        status="Ready",
        selected_for_generation=True,
    )


def _save_multistop_docx(
    *,
    record: BolMultistopRecord,
    bol_label: str,
    selected_facility: BolFacilityRecord,
    batch_comment: str | None,
    resolved_template: Path,
    output_root: Path,
    base_name: str,
    replacements: dict[str, str],
    document_type: str,
    stop_number: int | None,
    notices: list[DocxGenerationNotice],
) -> MultistopGeneratedDocxFile:
    doc = Document(str(resolved_template))
    _tighten_multistop_template_rows(doc)
    _replace_text_in_document(doc, replacements, include_xml_tree=True)
    _fit_combined_bol_number(doc, record.bol_number)
    _clean_combined_multistop_item_area(doc, record)
    for table in doc.tables:
        format_bol_item_detail_table(table)
    ship_from_populated = _populate_ship_from_block(doc, selected_facility)
    if not ship_from_populated:
        notices.append(
            DocxGenerationNotice(
                bol_number=bol_label,
                message="Could not confirm ship-from block location in template.",
            )
        )
    bill_to_populated = _populate_combined_bill_to_block(doc, record)
    if not bill_to_populated:
        notices.append(
            DocxGenerationNotice(
                bol_number=bol_label,
                message="Could not confirm combined Bill To block location in template.",
            )
        )

    destination = _unique_destination_path(output_root, base_name, ".docx")
    filename = destination.name
    doc.save(str(destination))

    resolved_comment = _resolve_comment_for_record(record.comments, batch_comment)
    comment_label_populated = _postprocess_standard_comments_in_saved_docx(
        destination,
        resolved_comment,
    )
    if resolved_comment and not comment_label_populated:
        notices.append(
            DocxGenerationNotice(
                bol_number=bol_label,
                message=(
                    "Resolved comment was non-empty but could not be confirmed "
                    "at the visible Comments label in word/document.xml."
                ),
            )
        )

    return MultistopGeneratedDocxFile(
        bol_number=bol_label,
        file_name=filename,
        file_path=str(destination.resolve()),
        document_type=document_type,
        load_number=record.load_number,
        stop_number=stop_number,
    )


def _save_individual_stop_docx(
    *,
    record: BolMultistopRecord,
    bol_label: str,
    selected_facility: BolFacilityRecord,
    batch_comment: str | None,
    resolved_template: Path,
    output_root: Path,
    base_name: str,
    stop_index: int,
    notices: list[DocxGenerationNotice],
) -> MultistopGeneratedDocxFile:
    stop = record.stops[stop_index]
    stop_record = _build_individual_stop_standard_record(record, stop_index)
    doc = Document(str(resolved_template))
    is_standard_template = resolved_template.name == STANDARD_TEMPLATE_PATH.name
    is_no_recourse_template = resolved_template.name == NO_RECOURSE_TEMPLATE_PATH.name
    resolved_comment = _resolve_comment_for_record(record.comments, batch_comment)
    record_notice_messages = _apply_standard_template_record_values(
        doc,
        stop_record,
        selected_facility,
        batch_comment,
        compact_standard_item_area=is_standard_template,
    )
    if is_standard_template:
        _clean_standard_individual_stop_item_area(doc, stop)
    elif is_no_recourse_template:
        _clean_no_recourse_individual_stop_item_area(doc, stop)
    for message in record_notice_messages:
        notices.append(DocxGenerationNotice(bol_number=bol_label, message=message))

    destination = _unique_destination_path(output_root, base_name, ".docx")
    filename = destination.name
    doc.save(str(destination))

    comment_label_populated = _postprocess_standard_comments_in_saved_docx(
        destination,
        resolved_comment,
    )
    if resolved_comment and not comment_label_populated:
        notices.append(
            DocxGenerationNotice(
                bol_number=bol_label,
                message=(
                    "Resolved comment was non-empty but could not be confirmed "
                    "at the visible Comments label in word/document.xml."
                ),
            )
        )

    return MultistopGeneratedDocxFile(
        bol_number=bol_label,
        file_name=filename,
        file_path=str(destination.resolve()),
        document_type="stop",
        load_number=record.load_number,
        stop_number=stop.stop_number,
    )


def generate_multistop_docx_set(
    records: list[BolMultistopRecord],
    selected_facility: BolFacilityRecord | None,
    batch_comment: str | None = None,
    template_path: Path | None = None,
    individual_stop_template_path: Path | None = None,
    output_dir: Path | None = None,
    file_name_prefix: str = "multistop_bol",
) -> StandardDocxGenerationResult:
    if selected_facility is None:
        raise ValueError(
            "No ship-from facility is selected. Select a facility in BOL Generator before DOCX generation."
        )

    resolved_template = template_path or MULTISTOP_TEMPLATE_PATH
    if not resolved_template.exists():
        raise FileNotFoundError(f"Template file not found: {resolved_template}")

    resolved_individual_stop_template = individual_stop_template_path or STANDARD_TEMPLATE_PATH
    if not resolved_individual_stop_template.exists():
        raise FileNotFoundError(
            f"Individual stop template file not found: {resolved_individual_stop_template}"
        )

    output_root = output_dir or Path(mkdtemp(prefix="kkg_multistop_bol_docx_"))
    output_root.mkdir(parents=True, exist_ok=True)

    generated: list[GeneratedDocxFile] = []
    skipped: list[SkippedDocxRecord] = []
    failed: list[FailedDocxRecord] = []
    notices: list[DocxGenerationNotice] = []

    for record in records:
        bol_label = record.bol_number or "(missing BOL #)"
        record.generation_skip_reason = None

        if not record.selected_for_generation:
            reason = "Record excluded in review."
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue

        if record.stop_count > 3:
            reason = "Unsupported stop count: more than 3 stops."
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue

        if not record.is_ready:
            reason = "Record is not ready for DOCX generation."
            if record.status == "Unsupported Stop Count":
                reason = "Unsupported stop count: more than 3 stops."
            elif record.missing_required_fields:
                reason = "Missing required data: " + ", ".join(record.missing_required_fields)
            elif record.issues:
                reason = "; ".join(record.issues)
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue

        try:
            safe_bol = _sanitize_filename_part(record.bol_number)
            safe_load = _sanitize_filename_part(record.load_number)
            combined_base_name = f"combined_multistop_bol_{safe_bol}_{safe_load}"

            generated.append(
                _save_multistop_docx(
                    record=record,
                    bol_label=bol_label,
                    selected_facility=selected_facility,
                    batch_comment=batch_comment,
                    resolved_template=resolved_template,
                    output_root=output_root,
                    base_name=combined_base_name,
                    replacements=_template_replacements(record),
                    document_type="combined",
                    stop_number=None,
                    notices=notices,
                )
            )

            for stop_index, stop in enumerate(record.stops):
                stop_base_name = f"stop_{stop.stop_number}_bol_{safe_bol}_{safe_load}"
                generated.append(
                    _save_individual_stop_docx(
                        record=record,
                        bol_label=bol_label,
                        selected_facility=selected_facility,
                        batch_comment=batch_comment,
                        resolved_template=resolved_individual_stop_template,
                        output_root=output_root,
                        base_name=stop_base_name,
                        stop_index=stop_index,
                        notices=notices,
                    )
                )
        except Exception as exc:
            failed.append(FailedDocxRecord(bol_number=bol_label, error=str(exc)))

    if not generated and not failed:
        raise ValueError("No selected and ready records are available for DOCX generation.")

    return StandardDocxGenerationResult(
        output_dir=str(output_root.resolve()),
        generated_files=generated,
        skipped_records=skipped,
        failed_records=failed,
        notices=notices,
    )
