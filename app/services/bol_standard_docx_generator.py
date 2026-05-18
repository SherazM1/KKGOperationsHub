"""DOCX generation service for Standard-family BOL records."""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from tempfile import mkdtemp
from xml.sax.saxutils import escape
import re
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

from app.models.bol_standard_record import BolStandardItemLine, BolStandardRecord
from app.utils.bol_facilities import BolFacilityRecord, facility_to_ship_from


STANDARD_TEMPLATE_PATH = Path("app/templates/standard_bol_template.docx")
NO_RECOURSE_TEMPLATE_PATH = Path("app/templates/no_recourse_bol_template.docx")
DEFAULT_TEMPLATE_PATH = STANDARD_TEMPLATE_PATH
LEFT_MERGE = "\u00ab"
RIGHT_MERGE = "\u00bb"
ITEM_TOKEN_ALIASES: dict[str, tuple[str, ...]] = {
    "QTY": ("QTY", "QTY_2", "QTY_3", "QTY_4"),
    "TYPE": ("TYPE", "TYPE_2", "TYPE_3", "TYPE_4"),
    "PO": ("PO_", "PO_2", "PO_3", "PO_4"),
    "ITEM_DESCRIPTION": (
        "Item_Description",
        "Item_2_Description",
        "Item_Description_3",
        "Item_4_Description",
    ),
    "ITEM_NUMBER": ("Item_Number", "Item_Number_2", "Item_Number_3", "Item_Number_4"),
    "UPC": ("UPC_", "UPC__2", "UPC__3", "UPC__4"),
    "WEIGHT": ("WEIGHT", "WEIGHT_2", "WEIGHT_3", "WEIGHT_4"),
}


def _tok(name: str) -> str:
    return f"{LEFT_MERGE}{name}{RIGHT_MERGE}"


ITEM_PLACEHOLDER_TOKENS: tuple[str, ...] = tuple(
    _tok(alias) for aliases in ITEM_TOKEN_ALIASES.values() for alias in aliases
)
DOCUMENT_FONT_NAME = "Arial"


def resolve_template_path_for_mode(mode: str) -> Path:
    if mode == "Standard":
        return STANDARD_TEMPLATE_PATH
    if mode == "No Recourse":
        return NO_RECOURSE_TEMPLATE_PATH
    raise ValueError(
        f"Unsupported BOL mode for Standard-family generation: {mode}. "
        "Use Standard or No Recourse."
    )


def resolve_output_filename_prefix_for_mode(mode: str) -> str:
    if mode == "Standard":
        return "standard_bol"
    if mode == "No Recourse":
        return "no_recourse_bol"
    raise ValueError(
        f"Unsupported BOL mode for Standard-family generation: {mode}. "
        "Use Standard or No Recourse."
    )


@dataclass(slots=True)
class GeneratedDocxFile:
    bol_number: str
    file_name: str
    file_path: str


@dataclass(slots=True)
class SkippedDocxRecord:
    bol_number: str
    reason: str


@dataclass(slots=True)
class FailedDocxRecord:
    bol_number: str
    error: str


@dataclass(slots=True)
class DocxGenerationNotice:
    bol_number: str
    message: str


@dataclass(slots=True)
class StandardDocxGenerationResult:
    output_dir: str
    generated_files: list[GeneratedDocxFile]
    skipped_records: list[SkippedDocxRecord]
    failed_records: list[FailedDocxRecord]
    notices: list[DocxGenerationNotice]

    @property
    def generated_count(self) -> int:
        return len(self.generated_files)

    @property
    def skipped_count(self) -> int:
        return len(self.skipped_records)

    @property
    def failed_count(self) -> int:
        return len(self.failed_records)


def _sanitize_filename_part(value: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in ("-", "_") else "_" for char in value)
    cleaned = cleaned.strip("_")
    return cleaned or "unknown"


def _normalize_bol_type(bol_type: str | None) -> str:
    normalized = (bol_type or "PLT").strip().upper()
    return normalized if normalized in {"PLT", "CASE"} else "PLT"


def _qty_type_header(qty_type: str | None) -> str:
    normalized = (qty_type or "PLT").strip().upper()
    return "Case Qty" if normalized == "CASE" else "Pallet Qty"


def _unique_destination_path(directory: Path, base_name: str, extension: str) -> Path:
    candidate = directory / f"{base_name}{extension}"
    if not candidate.exists():
        return candidate

    suffix = 2
    while True:
        candidate = directory / f"{base_name}_{suffix}{extension}"
        if not candidate.exists():
            return candidate
        suffix += 1


def _replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    text_nodes = paragraph._p.findall(".//w:t", paragraph._p.nsmap)
    instr_nodes = paragraph._p.findall(".//w:instrText", paragraph._p.nsmap)
    for node in [*text_nodes, *instr_nodes]:
        text = node.text or ""
        updated = text
        for source, target in replacements.items():
            if source in updated:
                updated = updated.replace(source, target)
        if updated != text:
            node.text = updated


def _replace_text_in_document(
    doc: Document, replacements: dict[str, str], *, include_xml_tree: bool = True
) -> None:
    def _replace_in_table_collection(tables: list[Table]) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, replacements)
                    _replace_in_table_collection(cell.tables)

    def _replace_in_element_tree(element) -> None:
        text_nodes = element.findall(".//w:t", element.nsmap)
        instr_nodes = element.findall(".//w:instrText", element.nsmap)
        for node in [*text_nodes, *instr_nodes]:
            text = node.text or ""
            updated = text
            for source, target in replacements.items():
                if source in updated:
                    updated = updated.replace(source, target)
            if updated != text:
                node.text = updated

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    _replace_in_table_collection(doc.tables)
    if include_xml_tree:
        _replace_in_element_tree(doc.element)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        _replace_in_table_collection(section.header.tables)
        if include_xml_tree:
            _replace_in_element_tree(section.header._element)

        for paragraph in section.footer.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        _replace_in_table_collection(section.footer.tables)
        if include_xml_tree:
            _replace_in_element_tree(section.footer._element)


def _replace_tokens_in_row_element(row_element, replacements: dict[str, str]) -> None:
    text_nodes = row_element.findall(".//w:t", row_element.nsmap)
    for node in text_nodes:
        text = node.text or ""
        for source, target in replacements.items():
            if source in text:
                text = text.replace(source, target)
        node.text = text


def _row_has_any_token_text(row, tokens: tuple[str, ...]) -> bool:
    row_text = " ".join(cell.text for cell in row.cells)
    return any(token in row_text for token in tokens)


def _document_contains_token(doc: Document, token: str) -> bool:
    return token in doc.element.xml


def _override_consignee_street(doc: Document, consignee_street: str) -> None:
    for table in doc.tables:
        in_consignee_block = False
        for row in table.rows:
            row_cells = row.cells
            if not row_cells:
                continue

            first_cell_text = row_cells[0].text.strip().upper()
            row_text = " ".join(cell.text.strip() for cell in row_cells).upper()

            if "TO (CONSIGNEE)" in row_text:
                in_consignee_block = True
                continue

            if in_consignee_block and first_cell_text == "STREET" and len(row_cells) > 1:
                row_cells[1].text = consignee_street
                return

            if in_consignee_block and "QTY" in row_text and "ITEM DESCRIPTION" in row_text:
                return


def _suppress_duplicate_ship_from_city_state_line(doc: Document, ship_from_location: str) -> None:
    location_value = ship_from_location.strip()
    if not location_value:
        return

    def _iter_tables(tables: list[Table]):
        for table in tables:
            yield table
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_tables(cell.tables)

    all_tables: list[Table] = list(_iter_tables(doc.tables))
    for section in doc.sections:
        all_tables.extend(list(_iter_tables(section.header.tables)))
        all_tables.extend(list(_iter_tables(section.footer.tables)))

    for table in all_tables:
        in_ship_from_block = False
        city_state_row_count = 0

        for row in table.rows:
            row_cells = row.cells
            if not row_cells:
                continue

            row_text_upper = " ".join(cell.text.strip() for cell in row_cells).upper()

            if "FROM (SHIPPER)" in row_text_upper:
                in_ship_from_block = True
                continue

            if in_ship_from_block and "TO (CONSIGNEE)" in row_text_upper:
                break

            if not in_ship_from_block:
                continue

            row_text = " ".join(cell.text.strip() for cell in row_cells)
            first_cell_label = row_cells[0].text.strip().upper().replace(" ", "")
            is_city_state_row = (
                "CITY/ST/ZIP" in row_text_upper
                or "CITY/STATE/ZIP" in row_text_upper
                or "CITY/ST/ZIP" in first_cell_label
                or "CITY/STATE/ZIP" in first_cell_label
            )
            has_location_value = location_value in row_text

            if not (is_city_state_row or has_location_value):
                continue

            city_state_row_count += 1
            if city_state_row_count > 1:
                for cell in row_cells:
                    cell.text = ""


def _item_row_replacements(line: BolStandardItemLine, bol_type: str | None = None) -> dict[str, str]:
    rendered_type = _normalize_bol_type(bol_type)
    replacements: dict[str, str] = {}
    for alias in ITEM_TOKEN_ALIASES["QTY"]:
        replacements[_tok(alias)] = line.pallet_qty
    for alias in ITEM_TOKEN_ALIASES["TYPE"]:
        replacements[_tok(alias)] = rendered_type
    for alias in ITEM_TOKEN_ALIASES["PO"]:
        replacements[_tok(alias)] = line.po_number
    for alias in ITEM_TOKEN_ALIASES["ITEM_DESCRIPTION"]:
        replacements[_tok(alias)] = line.item_description
    for alias in ITEM_TOKEN_ALIASES["ITEM_NUMBER"]:
        replacements[_tok(alias)] = line.item_number
    for alias in ITEM_TOKEN_ALIASES["UPC"]:
        replacements[_tok(alias)] = line.upc
    for alias in ITEM_TOKEN_ALIASES["WEIGHT"]:
        replacements[_tok(alias)] = line.weight_each
    return replacements


def _total_qty_display(total_skids: float) -> str:
    return str(int(total_skids)) if float(total_skids).is_integer() else str(total_skids)


def _parse_numeric(value: str) -> float | None:
    cleaned = (value or "").replace(",", "").strip()
    if not cleaned:
        return None
    numeric_match = re.search(r"-?\d+(?:\.\d+)?", cleaned)
    if numeric_match is None:
        return None
    try:
        return float(numeric_match.group(0))
    except ValueError:
        return None


def _format_number(value: float) -> str:
    return str(int(value)) if float(value).is_integer() else f"{value:.2f}".rstrip("0").rstrip(".")


def _set_run_font_name(run, font_name: str = DOCUMENT_FONT_NAME) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for attr_name in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr_name}"), font_name)
    for attr_name in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        theme_attr = qn(f"w:{attr_name}")
        if theme_attr in r_fonts.attrib:
            del r_fonts.attrib[theme_attr]


def _normalize_font_in_xml_tree(element, font_name: str = DOCUMENT_FONT_NAME) -> None:
    for r_pr in element.findall(".//w:rPr", element.nsmap):
        r_fonts = r_pr.find("w:rFonts", element.nsmap)
        if r_fonts is None:
            r_fonts = r_pr.makeelement(qn("w:rFonts"), nsmap=r_pr.nsmap)
            r_pr.insert(0, r_fonts)
        for attr_name in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attr_name}"), font_name)
        for attr_name in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            theme_attr = qn(f"w:{attr_name}")
            if theme_attr in r_fonts.attrib:
                del r_fonts.attrib[theme_attr]


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _normalize_document_font(doc: Document, font_name: str = DOCUMENT_FONT_NAME) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font_name(run, font_name)

    for table in doc.tables:
        for paragraph in _iter_table_paragraphs(table):
            for run in paragraph.runs:
                _set_run_font_name(run, font_name)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                _set_run_font_name(run, font_name)
        for table in section.header.tables:
            for paragraph in _iter_table_paragraphs(table):
                for run in paragraph.runs:
                    _set_run_font_name(run, font_name)

        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                _set_run_font_name(run, font_name)
        for table in section.footer.tables:
            for paragraph in _iter_table_paragraphs(table):
                for run in paragraph.runs:
                    _set_run_font_name(run, font_name)

    _normalize_font_in_xml_tree(doc.element, font_name)
    for section in doc.sections:
        _normalize_font_in_xml_tree(section.header._element, font_name)
        _normalize_font_in_xml_tree(section.footer._element, font_name)


def _set_cell_text_size(cell, size: Pt) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = size


def _set_cell_nowrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(cell._tc.makeelement(qn("w:noWrap"), nsmap=cell._tc.nsmap))


def _remove_cell_nowrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for no_wrap in list(tc_pr.findall(qn("w:noWrap"))):
        tc_pr.remove(no_wrap)


def _row_is_item_header(row) -> bool:
    row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip()).upper()
    return (
        "QTY" in row_text
        and "TYPE" in row_text
        and "PO #" in row_text
        and "ITEM DESCRIPTION" in row_text
    )


def _apply_header_fit_cleanup(doc: Document, record: BolStandardRecord) -> None:
    long_values = {
        value.strip()
        for value in (
            record.bol_number,
            record.po_number,
            record.kk_po_number,
            record.kk_load_number,
            record.pickup_number,
        )
        if len((value or "").strip()) >= 11
    }
    if not long_values:
        return

    header_labels = (
        "BOL #",
        "PO #",
        "CARRIER PRO #",
        "KK PO #",
        "KKG LOAD #",
        "KK LOAD #",
        "PICK UP #",
        "DELIVERY APPT.",
        "APPT #",
    )
    for table in doc.tables:
        for row in table.rows:
            if _row_is_item_header(row):
                break
            for index, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_text_upper = cell_text.upper()
                has_long_value = any(value in cell_text for value in long_values)
                has_header_label = any(label in cell_text_upper for label in header_labels)
                if not cell_text or not has_long_value:
                    continue
                _set_cell_text_size(cell, Pt(7.5))
                if has_header_label and index + 1 < len(row.cells):
                    value_cell = row.cells[index + 1]
                    _set_cell_text_size(value_cell, Pt(7.5))
                    _set_cell_nowrap(value_cell)
                else:
                    _set_cell_nowrap(cell)


def _apply_type_column_fit_cleanup(doc: Document) -> None:
    for table in doc.tables:
        header_idx = None
        type_col_indexes: list[int] = []
        for idx, row in enumerate(table.rows):
            row_text_upper = " ".join(cell.text.strip() for cell in row.cells).upper()
            if (
                "TYPE" in row_text_upper
                and "PO #" in row_text_upper
                and "ITEM DESCRIPTION" in row_text_upper
            ):
                header_idx = idx
                type_col_indexes = [
                    col_idx
                    for col_idx, cell in enumerate(row.cells)
                    if cell.text.strip().upper() == "TYPE"
                ]
                break
        if header_idx is None or not type_col_indexes:
            continue

        for row in table.rows[header_idx + 1 :]:
            row_text_upper = " ".join(cell.text.strip() for cell in row.cells).upper()
            if "TOTALS" in row_text_upper:
                break
            for col_idx in type_col_indexes:
                if col_idx >= len(row.cells):
                    continue
                cell = row.cells[col_idx]
                if cell.text.strip().upper() in {"PLT", "CASE"}:
                    _set_cell_text_size(cell, Pt(7))
                    _set_cell_nowrap(cell)


def _apply_subject7_fit_cleanup(doc: Document) -> None:
    for paragraph in doc.element.findall(".//w:p", doc.element.nsmap):
        paragraph_text = "".join(
            node.text or "" for node in paragraph.findall(".//w:t", paragraph.nsmap)
        )
        if "SUBJECT TO SECTION 7" not in paragraph_text.upper():
            continue
        for r_pr in paragraph.findall(".//w:rPr", paragraph.nsmap):
            sz = r_pr.find("w:sz", paragraph.nsmap)
            if sz is None:
                sz = r_pr.makeelement(qn("w:sz"), nsmap=r_pr.nsmap)
                r_pr.append(sz)
            sz.set(qn("w:val"), "14")

            sz_cs = r_pr.find("w:szCs", paragraph.nsmap)
            if sz_cs is None:
                sz_cs = r_pr.makeelement(qn("w:szCs"), nsmap=r_pr.nsmap)
                r_pr.append(sz_cs)
            sz_cs.set(qn("w:val"), "14")

    for body_pr in doc.element.findall(".//wps:bodyPr", doc.element.nsmap):
        body_pr.set("horzOverflow", "clip")
        body_pr.set("vertOverflow", "clip")


def _item_line_has_data(line: BolStandardItemLine) -> bool:
    return any(
        value.strip()
        for value in (
            line.pallet_qty,
            line.po_number,
            line.item_description,
            line.item_number,
            line.upc,
            line.skids,
            line.weight_each,
            line.total_weight,
        )
    )


def _format_ship_date_for_template(raw_ship_date: str) -> str:
    value = (raw_ship_date or "").strip()
    if not value:
        return ""

    normalized = value.replace("T", " ")
    for sep in (" ", "."):
        if sep in normalized:
            date_candidate = normalized.split(sep, 1)[0].strip()
            if date_candidate and any(char.isdigit() for char in date_candidate):
                normalized = date_candidate
                break

    parse_formats = (
        "%Y-%m-%d",
        "%m/%d/%Y",
        "%m/%d/%y",
        "%m-%d-%Y",
        "%m-%d-%y",
    )
    for parse_format in parse_formats:
        try:
            parsed = datetime.strptime(normalized, parse_format)
            return parsed.strftime("%m/%d/%Y")
        except ValueError:
            continue

    try:
        parsed_iso = datetime.fromisoformat(value.replace("Z", ""))
        return parsed_iso.strftime("%m/%d/%Y")
    except ValueError:
        return normalized


def _populate_item_table(
    table: Table,
    item_lines: list[BolStandardItemLine],
    total_qty: float,
    *,
    bol_type: str | None = None,
    qty_type: str = "PLT",
    compact_standard_item_area: bool = False,
    filter_blank_item_lines: bool = False,
) -> None:
    rendered_type = _normalize_bol_type(bol_type)
    rendered_qty_header = _qty_type_header(qty_type)
    rendered_item_lines = (
        [line for line in item_lines if _item_line_has_data(line)]
        if filter_blank_item_lines
        else item_lines
    )

    header_idx = None
    item_row_indices: list[int] = []
    total_qty_idx = None
    totals_label_idx = None

    for idx, row in enumerate(table.rows):
        row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
        row_text_upper = row_text.upper()

        if (
            header_idx is None
            and "QTY" in row_text_upper
            and "TYPE" in row_text_upper
            and "PO #" in row_text_upper
            and "ITEM DESCRIPTION" in row_text_upper
        ):
            header_idx = idx
            continue

        if header_idx is not None and idx > header_idx and _row_has_any_token_text(
            row, ITEM_PLACEHOLDER_TOKENS
        ):
            item_row_indices.append(idx)

        if header_idx is not None and idx > header_idx and _tok("TOTAL_QTY") in row_text:
            total_qty_idx = idx
        if header_idx is not None and idx > header_idx and "TOTALS" in row_text_upper:
            if totals_label_idx is None:
                totals_label_idx = idx

    if header_idx is None:
        raise ValueError("Could not locate the item-table header in the DOCX template.")
    if not item_row_indices:
        raise ValueError("Could not locate item template rows in the DOCX template.")
    if total_qty_idx is None:
        raise ValueError("Could not locate the TOTAL_QTY row in the DOCX template.")

    first_idx = item_row_indices[0]
    contiguous_item_row_indices = [first_idx]
    for idx in item_row_indices[1:]:
        if idx == contiguous_item_row_indices[-1] + 1:
            contiguous_item_row_indices.append(idx)
        else:
            break

    insertion_anchor_idx = total_qty_idx
    if compact_standard_item_area and totals_label_idx is not None:
        insertion_anchor_idx = totals_label_idx

    table_xml = table._tbl
    template_trs = [deepcopy(table.rows[idx]._tr) for idx in contiguous_item_row_indices]
    anchor_tr = table.rows[insertion_anchor_idx]._tr
    data_row_start_idx = contiguous_item_row_indices[0]

    if compact_standard_item_area:
        remove_start_idx = contiguous_item_row_indices[0]
        remove_end_idx = insertion_anchor_idx - 1
        rows_to_remove = list(range(remove_start_idx, remove_end_idx + 1))
        for idx in sorted(rows_to_remove, reverse=True):
            table_xml.remove(table.rows[idx]._tr)

        for idx, line in enumerate(rendered_item_lines):
            new_tr = deepcopy(template_trs[idx % len(template_trs)])
            _replace_tokens_in_row_element(new_tr, _item_row_replacements(line, rendered_type))
            anchor_tr.addprevious(new_tr)
    else:
        empty_item_replacements = {token: "" for token in ITEM_PLACEHOLDER_TOKENS}
        for template_idx, row_idx in enumerate(contiguous_item_row_indices):
            row_tr = table.rows[row_idx]._tr
            if template_idx < len(rendered_item_lines):
                _replace_tokens_in_row_element(
                    row_tr, _item_row_replacements(rendered_item_lines[template_idx], rendered_type)
                )
            else:
                _replace_tokens_in_row_element(row_tr, empty_item_replacements)

        if len(rendered_item_lines) > len(contiguous_item_row_indices):
            for idx, line in enumerate(
                rendered_item_lines[len(contiguous_item_row_indices):], start=0
            ):
                new_tr = deepcopy(template_trs[idx % len(template_trs)])
                _replace_tokens_in_row_element(new_tr, _item_row_replacements(line, rendered_type))
                anchor_tr.addprevious(new_tr)

    header_cells = [cell.text.strip().upper() for cell in table.rows[header_idx].cells]
    qty_col_indexes = [
        idx for idx, cell_text in enumerate(header_cells)
        if "QTY" in cell_text and "# SKIDS" not in cell_text
    ]
    if qty_col_indexes:
        table.rows[header_idx].cells[qty_col_indexes[0]].text = rendered_qty_header

    skids_col_indexes = [idx for idx, cell_text in enumerate(header_cells) if "# SKIDS" in cell_text]
    type_col_indexes = [idx for idx, cell_text in enumerate(header_cells) if cell_text == "TYPE"]
    weight_col_indexes = [idx for idx, cell_text in enumerate(header_cells) if "WEIGHT" in cell_text]

    totals_anchor_idx = None
    for idx in range(header_idx + 1, len(table.rows)):
        row_text_upper = " ".join(cell.text.strip() for cell in table.rows[idx].cells).upper()
        if "TOTALS" in row_text_upper:
            totals_anchor_idx = idx
            break

    if totals_anchor_idx is None:
        raise ValueError("Could not locate the TOTALS row in the DOCX template.")

    if compact_standard_item_area:
        item_start_idx = totals_anchor_idx - len(rendered_item_lines)
        for line_offset, line in enumerate(rendered_item_lines):
            row_idx = item_start_idx + line_offset
            if row_idx <= header_idx or row_idx >= len(table.rows):
                continue

            row_cells = table.rows[row_idx].cells
            for col_idx in qty_col_indexes:
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = line.pallet_qty
            for col_idx in skids_col_indexes:
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = line.skids
            for col_idx in type_col_indexes:
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = rendered_type
            for col_idx in weight_col_indexes:
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = line.total_weight or line.weight_each

    total_pallet_qty_value = 0.0
    total_skids_value = 0.0
    total_weight_value = 0.0
    has_pallet_qty_value = False
    use_line_total_weight = any((line.total_weight or "").strip() for line in rendered_item_lines)
    has_total_weight_value = False
    for line in rendered_item_lines:
        numeric_pallet_qty = _parse_numeric(line.pallet_qty)
        if numeric_pallet_qty is not None:
            total_pallet_qty_value += numeric_pallet_qty
            has_pallet_qty_value = True

        numeric_skids = _parse_numeric(line.skids)
        if numeric_skids is not None:
            total_skids_value += numeric_skids

        weight_source = line.total_weight if use_line_total_weight else line.weight_each
        numeric_weight = _parse_numeric(weight_source)
        if numeric_weight is not None:
            total_weight_value += numeric_weight
            if use_line_total_weight:
                has_total_weight_value = True

    if use_line_total_weight and not has_total_weight_value:
        for line in rendered_item_lines:
            numeric_weight = _parse_numeric(line.weight_each)
            if numeric_weight is not None:
                total_weight_value += numeric_weight

    total_qty_display = (
        _format_number(total_pallet_qty_value)
        if filter_blank_item_lines and has_pallet_qty_value
        else _total_qty_display(total_qty)
    )
    total_skids_display = _format_number(total_skids_value)
    total_weight_display = _format_number(total_weight_value)

    if compact_standard_item_area:
        totals_row_cells = table.rows[totals_anchor_idx].cells
        for col_idx in qty_col_indexes:
            if col_idx < len(totals_row_cells):
                totals_row_cells[col_idx].text = total_qty_display
        for col_idx in skids_col_indexes:
            if col_idx < len(totals_row_cells):
                totals_row_cells[col_idx].text = total_skids_display
        for col_idx in weight_col_indexes:
            if col_idx < len(totals_row_cells):
                totals_row_cells[col_idx].text = total_weight_display

        totals_replacements = {
            _tok("TOTAL_QTY"): "",
            _tok("TOTAL_WEIGHT"): "",
        }
        for row in table.rows:
            _replace_tokens_in_row_element(row._tr, totals_replacements)
    else:
        totals_replacements = {
            _tok("TOTAL_QTY"): total_qty_display,
            _tok("TOTAL_WEIGHT"): total_weight_display,
        }
        for row in table.rows:
            _replace_tokens_in_row_element(row._tr, totals_replacements)


def _resolve_comment_for_record(record_comment: str, batch_comment: str | None) -> str:
    record_value = (record_comment or "").strip()
    if record_value:
        return record_value

    return (batch_comment or "").strip()


def _postprocess_comments_in_saved_docx(destination: Path, resolved_comment: str) -> bool:
    xml_path = "word/document.xml"
    with zipfile.ZipFile(destination, "r") as archive:
        if xml_path not in archive.namelist():
            return False
        file_payloads = {name: archive.read(name) for name in archive.namelist()}

    xml_text = file_payloads[xml_path].decode("utf-8", errors="ignore")
    updated_xml, comment_label_populated = _postprocess_comments_in_document_xml(
        xml_text,
        resolved_comment,
    )

    if updated_xml == xml_text:
        return comment_label_populated

    file_payloads[xml_path] = updated_xml.encode("utf-8")
    with zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, payload in file_payloads.items():
            archive.writestr(name, payload)
    return comment_label_populated


def _strip_comment_mergefield_shapes(xml_text: str) -> str:
    updated = xml_text
    previous = None
    while previous != updated:
        previous = updated
        updated = re.sub(
            r"<mc:AlternateContent\b(?:(?!</mc:AlternateContent>).)*?"
            r"MERGEFIELD\s+COMMENTS"
            r"(?:(?!</mc:AlternateContent>).)*?</mc:AlternateContent>",
            "",
            updated,
            flags=re.DOTALL,
        )
        updated = re.sub(
            r"<w:p\b(?:(?!</w:p>).)*?"
            r"MERGEFIELD\s+COMMENTS"
            r"(?:(?!</w:p>).)*?</w:p>",
            "",
            updated,
            flags=re.DOTALL,
        )
    return updated


def _clear_comment_placeholders(xml_text: str) -> str:
    placeholder_patterns = (
        r"\s*MERGEFIELD\s+COMMENTS\s*",
        re.escape(_tok("COMMENTS")),
        re.escape("\u00ab COMMENTS \u00bb"),
        re.escape("\u00a0\u00abCOMMENTS\u00bb"),
        r"&lt;&lt;\s*COMMENTS\s*&gt;&gt;",
        r"<<\s*COMMENTS\s*>>",
    )
    updated = xml_text
    for pattern in placeholder_patterns:
        updated = re.sub(pattern, "", updated, flags=re.IGNORECASE)
    return updated


def _clear_existing_comment_label_values(xml_text: str) -> str:
    label_pattern = re.compile(
        r"(<w:t(?:\s+[^>]*)?>)(COMMENTS?:)(?:\s*[^<]*)(</w:t>)",
        flags=re.IGNORECASE,
    )
    return label_pattern.sub(lambda match: f"{match.group(1)}{match.group(2)}{match.group(3)}", xml_text)


def _populate_first_comment_label(xml_text: str, resolved_comment: str) -> tuple[str, bool]:
    if not resolved_comment:
        return xml_text, False

    safe_comment = escape(resolved_comment)
    label_pattern = re.compile(
        r"(<w:t(?:\s+[^>]*)?>)(COMMENTS?:)(?:\s*[^<]*)(</w:t>)",
        flags=re.IGNORECASE,
    )
    alternate_content_pattern = re.compile(
        r"<mc:AlternateContent\b(?:(?!</mc:AlternateContent>).)*?</mc:AlternateContent>",
        flags=re.DOTALL,
    )

    def _replace_label(match: re.Match[str]) -> str:
        return f"{match.group(1)}{match.group(2)} {safe_comment}{match.group(3)}"

    for alternate_match in alternate_content_pattern.finditer(xml_text):
        alternate_content = alternate_match.group(0)
        if not label_pattern.search(alternate_content):
            continue

        updated_alternate_content = label_pattern.sub(_replace_label, alternate_content)
        return (
            xml_text[: alternate_match.start()]
            + updated_alternate_content
            + xml_text[alternate_match.end() :],
            True,
        )

    updated_xml, replacement_count = label_pattern.subn(_replace_label, xml_text, count=1)
    return updated_xml, replacement_count > 0


def _postprocess_comments_in_document_xml(
    xml_text: str,
    resolved_comment: str,
) -> tuple[str, bool]:
    updated_xml = _strip_comment_mergefield_shapes(xml_text)
    updated_xml = _clear_comment_placeholders(updated_xml)
    updated_xml = _clear_existing_comment_label_values(updated_xml)
    updated_xml, populated = _populate_first_comment_label(updated_xml, resolved_comment.strip())
    return updated_xml, populated


def _apply_template_record_values(
    doc: Document,
    record: BolStandardRecord,
    selected_facility: BolFacilityRecord,
    batch_comment: str | None,
    *,
    bol_type: str | None = None,
    qty_type: str = "PLT",
    compact_standard_item_area: bool = False,
    filter_blank_item_lines: bool = False,
) -> list[str]:
    notices: list[str] = []
    has_explicit_pickup_token = _document_contains_token(doc, _tok("Pick_Up_"))
    replacements = {
        _tok("BOL"): record.bol_number,
        _tok("SHIP_DATE"): _format_ship_date_for_template(record.ship_date),
        _tok("CARRIER"): record.carrier,
        _tok("Carrier_Pro_"): record.carrier_pro_number,
        _tok("HOST_PO"): record.po_number,
        _tok("KKG_PO"): record.kk_po_number,
        _tok("KKG_LOAD_"): record.kk_load_number,
        _tok("Pick_Up_"): record.pickup_number,
        _tok("TRACKER_"): "",
        # Suppress competing mergefield comment box; visible "Comments:" label is authoritative.
        _tok("COMMENTS"): "",
        _tok("SHIP_FROM"): record.ship_from.company,
        _tok("SHIP_FROM_ADDRESS"): record.ship_from.street,
        _tok("SHIP_FROM_CITY_STATE_ZIP"): "" if has_explicit_pickup_token else record.pickup_number,
        _tok("SHIP_TO_NAME"): record.consignee_company,
        _tok("SHIP_TO_ADDRESS"): record.consignee_street,
        _tok("SHIP_TO_CITY_STATE_ZIP"): record.consignee_city_state_zip,
        _tok("DC"): record.dc_number,
        _tok("BILL_TO"): record.bill_to.company,
        _tok("BILL_TO_ADDRESS"): record.bill_to.street,
        _tok("BILL_TO_CITY_SATE_ZIP"): record.bill_to.city_state_zip,
    }
    _replace_text_in_document(
        doc,
        replacements,
        include_xml_tree=compact_standard_item_area,
    )
    _replace_text_in_document(
        doc,
        {
            "<<COMMENTS>>": "",
            "<< COMMENTS >>": "",
            "\u00ab COMMENTS \u00bb": "",
            "\u00a0\u00abCOMMENTS\u00bb": "",
            " MERGEFIELD COMMENTS ": "",
            "MERGEFIELD COMMENTS": "",
        },
        include_xml_tree=compact_standard_item_area,
    )
    _suppress_duplicate_ship_from_city_state_line(doc, record.ship_from.city_state_zip)
    _override_consignee_street(doc, record.consignee_street)

    last_error: Exception | None = None
    for table in doc.tables:
        try:
            _populate_item_table(
                table,
                record.item_lines,
                record.total_skids,
                bol_type=bol_type,
                qty_type=qty_type,
                compact_standard_item_area=compact_standard_item_area,
                filter_blank_item_lines=filter_blank_item_lines,
            )
            _apply_header_fit_cleanup(doc, record)
            _apply_type_column_fit_cleanup(doc)
            _apply_subject7_fit_cleanup(doc)
            _normalize_document_font(doc)
            return notices
        except ValueError as exc:
            last_error = exc

    if last_error is None:
        raise ValueError("Could not locate the item table in the DOCX template.")
    raise ValueError(f"Could not populate item table: {last_error}")


def generate_standard_docx_set(
    records: list[BolStandardRecord],
    selected_facility: BolFacilityRecord | None,
    batch_comment: str | None = None,
    bol_type: str | None = None,
    qty_type: str = "PLT",
    template_path: Path | None = None,
    output_dir: Path | None = None,
    file_name_prefix: str = "standard_bol",
) -> StandardDocxGenerationResult:
    if selected_facility is None:
        raise ValueError(
            "No ship-from facility is selected. Select a facility in BOL Generator before DOCX generation."
        )
    selected_ship_from = facility_to_ship_from(selected_facility)
    for record in records:
        record.ship_from = selected_ship_from

    resolved_template = template_path or DEFAULT_TEMPLATE_PATH
    if not resolved_template.exists():
        raise FileNotFoundError(f"Template file not found: {resolved_template}")

    output_root = output_dir or Path(mkdtemp(prefix="kkg_standard_bol_docx_"))
    output_root.mkdir(parents=True, exist_ok=True)

    generated: list[GeneratedDocxFile] = []
    skipped: list[SkippedDocxRecord] = []
    failed: list[FailedDocxRecord] = []
    notices: list[DocxGenerationNotice] = []

    for record in records:
        bol_label = record.bol_number or "(missing BOL #)"
        record.generation_skip_reason = None
        if not record.selected_for_generation:
            reason = "Record excluded in review."
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue
        if not record.is_ready:
            reason = "Record is not ready for DOCX generation."
            if record.missing_required_fields:
                reason = "Missing required data: " + ", ".join(record.missing_required_fields)
            elif record.issues:
                reason = "; ".join(record.issues)
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue

        try:
            doc = Document(str(resolved_template))
            is_standard_template = resolved_template.name == STANDARD_TEMPLATE_PATH.name
            is_no_recourse_template = resolved_template.name == NO_RECOURSE_TEMPLATE_PATH.name
            resolved_comment = _resolve_comment_for_record(record.comments, batch_comment)
            record_notices = _apply_template_record_values(
                doc,
                record,
                selected_facility,
                batch_comment,
                bol_type=bol_type,
                qty_type=qty_type,
                compact_standard_item_area=is_standard_template or is_no_recourse_template,
                filter_blank_item_lines=is_no_recourse_template,
            )

            safe_bol = _sanitize_filename_part(record.bol_number)
            destination = _unique_destination_path(
                output_root, f"{file_name_prefix}_{safe_bol}", ".docx"
            )
            filename = destination.name
            doc.save(str(destination))
            comment_label_populated = _postprocess_comments_in_saved_docx(
                destination, resolved_comment
            )
            if resolved_comment and not comment_label_populated:
                notices.append(
                    DocxGenerationNotice(
                        bol_number=bol_label,
                        message=(
                            "Debug: resolved comment was non-empty but could not be confirmed "
                            "at the visible Comments label in word/document.xml."
                        ),
                    )
                )

            generated.append(
                GeneratedDocxFile(
                    bol_number=bol_label,
                    file_name=filename,
                    file_path=str(destination.resolve()),
                )
            )
            for message in record_notices:
                notices.append(DocxGenerationNotice(bol_number=bol_label, message=message))
        except Exception as exc:
            failed.append(FailedDocxRecord(bol_number=bol_label, error=str(exc)))

    if not generated and not failed:
        raise ValueError("No selected and ready records are available for DOCX generation.")

    return StandardDocxGenerationResult(
        output_dir=str(output_root.resolve()),
        generated_files=generated,
        skipped_records=skipped,
        failed_records=failed,
        notices=notices,
    )
